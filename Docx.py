import PyPDF2
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re 
import requests
from bs4 import BeautifulSoup

from Excel import Excel


class Docx:
    def __init__(self) -> None:
        pdf_name = input("Write the pdf file name \n ->")
        self.pdf = open(f'analysis/{pdf_name}.pdf', 'rb')
        self.pdf_reader = PyPDF2.PdfReader(self.pdf)

        self.document = docx.Document()

        self.excel = Excel()
        self.excel.init()


        style = self.document.styles.add_style(
            'text_style', 
            WD_STYLE_TYPE.PARAGRAPH
        )
        style.font.name = 'Arial'
        style.font.size = Pt(10)

        style = self.document.styles.add_style(
            'subtitle_style', 
            WD_STYLE_TYPE.PARAGRAPH
        )
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.font.bold = True

        style = self.document.styles.add_style(
            'title_style', 
            WD_STYLE_TYPE.PARAGRAPH
        )
        style.font.name = 'Arial'
        style.font.size = Pt(16)
        style.font.bold = True

        button_style = self.document.styles.add_style(
            'button', 
            WD_STYLE_TYPE.PARAGRAPH
        )
        button_style.font.color.rgb = docx.shared.RGBColor(0, 176, 80)
        button_style.font.name = 'Arial'
        button_style.font.size = Pt(10)


        print("creating docx...")
        self.get_pdf_text()


    def get_pdf_text(self):
        text1 = "[quant] new stocks made [month]'s Exec Comp Aligned with ROIC Model Portfolio, available to members as of [publication_date]."
        text2 = "Our [title] Model Portfolio [percentage1] [status] the S&P 500 [percentage2] [analysis_dates] The best-performing stock in the portfolio was up [performing_percentage]. Overall, [stock_count] out [total_stock_count] [title] stocks outperformed the S&P 500 [analysis_dates]"


        self.text1_data = {}
        self.text2_data = {}
        
        page1_text = self.pdf_reader.pages[0].extract_text()
        important_data = self.get_data(page1_text, "Stocks", ".")


        publication_date = self.get_data(page1_text, "MONTHLY UPDATE", "Model")[0].strip()
        publication_day = int((self.get_data(page1_text, "MONTHLY UPDATE", "Model")[0].strip()).split("/")[1]) + 1
        self.text1_data['publication_date'] = f"{publication_date.split('/')[0]}/{publication_day}/{publication_date.split('/')[2]}"

        self.text1_data['quant'] = important_data[0].split()[4].capitalize()

        
        self.text2_data['title'] = self.get_data(page1_text, "Model Portfolio:", "ROIC")[0].strip()
        self.text2_data['percentage1'] = important_data[1].split()[4] 
        self.text2_data['percentage2'] = important_data[1].split()[9] + important_data[1].split()[10]


        if self.text2_data['percentage1'] > self.text2_data['percentage2']:
            self.text2_data['status'] = "outperformed"

        elif self.text2_data['percentage2'] > self.text2_data['percentage1']:
            self.text2_data['status'] = "underperformed"

        else:
            self.text2_data['status'] = "matched"
            


        self.text2_data['analysis_dates'] = important_data[2].strip()


        self.text2_data['total_stock_count'] = self.get_stock_count(1)
        self.text2_data['total_stock_count'] += self.get_stock_count(2)
        self.text2_data['total_stock_count'] += self.get_stock_count(3)


        page6_text = self.pdf_reader.pages[5].extract_text()
        important_data = self.get_data(page6_text, "From", "%")

        performing_percentage = important_data[1].split()[-1]
        self.text2_data['performing_percentage'] = f"{round(float(performing_percentage[0:-1]))}%"

        stock_count = 0

        for stock in important_data[1:]:
            if not "S&P" in stock:
                stock_count += 1
            else:
                self.text2_data['stock_count'] = stock_count
                break
    
        self.text1_data['month'] = self.text2_data['analysis_dates'].split()[5]
        

        text1 = self.edit_text(self.text1_data, text1)
        text2 = self.edit_text(self.text2_data, text2)

        self.write_docx(f"Featured Stock in {self.text1_data['month']}'s {self.text2_data['title']} Portfolio", 'title_style')        
        self.write_docx(text1, 'text_style')

        self.write_docx(f"Recap from {self.text1_data['month']} - {self.text2_data['analysis_dates'].split()[1]}'s Picks", 'subtitle_style')
        self.write_docx(text2, 'text_style')


        print("pdf data saved!")
        self.copy_texts()


    def copy_texts(self):
        button = self.document.add_paragraph()
        button.style = 'button'
        button.alignment = WD_ALIGN_PARAGRAPH.CENTER
        button.add_run('Buy the Exec Comp Aligned with ROIC Model Portfolio')


        paragraph1 = self.document.add_paragraph(
            "This report leverages our cutting-edge ", 
            'text_style'
        )
        self.create_hyperlink(
            paragraph1, 
            "Robo-Analyst technology", 
            'https://www.newconstructs.com/landing/robo-analyst-technology/'
        )
        paragraph1.add_run(" to deliver ")
        self.create_hyperlink(
            paragraph1, 
            "proven-superior", 
            'https://www.newconstructs.com/proof-of-the-superiority-of-our-data-models-ratings/'
        )
        paragraph1.add_run(" fundamental research and support more cost-effective fulfillment of the ")
        self.create_hyperlink(
            paragraph1, 
            "fiduciary duty of care.", 
            'https://www.newconstructs.com/even-without-the-law-fiduciary-rule-awareness-remains/'
        )
        

        paragraph2 = self.document.add_paragraph(
            "This Model Portfolio includes stocks that earn an ", 
            'text_style'
        )
        self.create_hyperlink(
            paragraph2, 
            "Attractive or Very Attractive", 
            'https://www.newconstructs.com/stock-rating-system/'
        )
        paragraph2.add_run(" rating and align executive compensation with improving ROIC. This combination provides a unique list of long ideas as the ")
        self.create_hyperlink(
            paragraph2, 
            "primary driver of shareholder value creation", 
            'https://www.newconstructs.com/roic-paradigm-linking-corporate-performance-valuation/'
        )
        paragraph2.add_run(" is return on invested capital (")
        self.create_hyperlink(
            paragraph2, 
            "ROIC", 
            'https://www.newconstructs.com/education-return-on-invested-capital/'
        )
        paragraph2.add_run(").")


        print("Static data saved!")
        self.get_excel_data()


    def get_excel_data(self):
        print("getting excel data...")


        text1_data = {}
        text1 = "New Feature Stock for [month]: [company] ([ticker]: $[price]/share)"


        text1_data['month'] = self.text1_data['month']


        text1_data['ticker'] = self.excel.reading_company_file()
        text1_data['company'], text1_data['price'] = self.get_company_data(text1_data['ticker'])


        text1 = self.edit_text(text1_data, text1)
        self.write_docx(text1, 'subtitle_style')



        self.excel.stock_file_manipulation()



        text2_data = {}
        text2 = "[company_name] has grown revenue and NOPAT by [revenue_5y]% and [nopat_5y]% compounded annually, respectively, since [5_years_back]. The company's NOPAT margin rose from [smaller_number]% in [smaller_number_year] to [current_percentage]% in 2022, while invested capital turns rose from [first_investment] to [last_investment] over the same time. "

        text2_data = self.excel.stock_file_data()  
        
        text2_data = self.excel.feature_stock_data(text2_data)
        
        text2_data['company_name'] = text1_data['company']
        text2_data['smaller_number'] = round(float(text2_data['smaller_number']))
        

        if text2_data['smaller_number'] < text2_data['current_percentage'] and text2_data['first_investment'] < text2_data['last_investment']:
            text2 += "Rising NOPAT margins and invested capital turns drive the company's return on invested capital ("

        elif text2_data['smaller_number'] < text2_data['current_percentage']:
            text2 += "Rising NOPAT margins drive the company's return on invested capital ("

        elif text2_data['first_investment'] < text2_data['last_investment']:
            text2 += "Invested capital turns drive the company's return on invested capital ("

        else:
            text2 += "The company's return on invested capital ("



        text2 = self.edit_text(text2_data, text2)
        paragraph = self.write_docx(text2, 'text_style')
        self.create_hyperlink(
            paragraph, 
            "ROIC", 
            'https://www.newconstructs.com/education-return-on-invested-capital/'
        )


        text2 = ") from [roic_first]% in [smaller_number_year] to [roic_last]% in 2022."
        
        text2 = self.edit_text(text2_data, text2)
        paragraph.add_run(text2)
        

        print("excel data saved!")
        self.save(self.text2_data['title'])


    def get_stock_count(self, page):
        page_text = self.pdf_reader.pages[page].extract_text()
        important_data = self.get_data(page_text, "Statement", "Statement")

        return len(important_data)


    def get_company_data(self, ticker):
        url = f'https://www.google.com/finance/quote/{ticker}'

        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')

        link = ((str(soup.select_one(f'a[href*="{ticker}:"]')).split()[1]).split("/")[2])[0:-1]


        url = f'https://www.google.com/finance/quote/{link}'
        
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')

        name = soup.find('div', {'class': ['zzDege']}).text

        price = round(float((str(soup.find('div', {'class': ['YMlKec fxKbKc']}).text)[1:]).replace(",", "")))


        return name, price


    def create_hyperlink(self, paragraph, sentence, link):
        part = paragraph.part
        r_id = part.relate_to(
            link, 
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, 
            is_external=True
        )

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        new_run.append(rPr)
        new_run.text = sentence
        hyperlink.append(new_run)

        r = paragraph.add_run ()
        r._r.append (hyperlink)
        r.font.underline = True


        return hyperlink


    def get_data(self, text, start, end):
        match = re.search(fr'{start} (.*{end})', text, re.DOTALL)

        try:
            data = match.group(1).split("\n")
            data.pop(0)

            return data
        
        except:
            print("No data found")

    
    def edit_text(self, new_words, text):
        new_text = text

        for word in new_words.items():
            new_text = new_text.replace(f"[{word[0]}]", f"{word[1]}")

        return new_text


    def write_docx(self, text, style):
        paragraph = self.document.add_paragraph(text, style)

        return paragraph


    def save(self, name):
        print("File saved!")

        self.document.save(f'created/{name}.docx')
        self.pdf.close()
    


a = Docx()
